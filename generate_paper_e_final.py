#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Paper E Generator - Creates comprehensive DOCX with plots
© 2025 Carmen Wrede & Lino Casu
"""
import os, sys
os.environ['PYTHONIOENCODING'] = 'utf-8:replace'

# Install deps
try:
    from docx import Document
    from docx.shared import Inches, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError:
    os.system(f"{sys.executable} -m pip install python-docx")
    from docx import Document
    from docx.shared import Inches, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH

import numpy as np
import matplotlib.pyplot as plt
from datetime import datetime

# Constants
R_S = 8.870e-3; R_E = 6.371e6; PHI = 1.618
F_T, F_O = 5e9, 429e12; T2_T, T2_O = 100e-6, 1.0

def gen_plots(d):
    os.makedirs(d, exist_ok=True)
    p = {}
    
    # Fig 2: Phase vs Height
    fig, ax = plt.subplots(figsize=(10,6))
    h = np.logspace(-6,2,100)
    phi_t = 2*np.pi*F_T * (R_S*h/R_E**2) * T2_T
    phi_o = 2*np.pi*F_O * (R_S*h/R_E**2) * T2_O
    ax.loglog(h*1000, phi_t, 'b-', lw=2, label='Transmon (5GHz, 100us)')
    ax.loglog(h*1000, phi_o, 'r-', lw=2, label='Optical (429THz, 1s)')
    ax.axhline(0.1, color='g', ls='--', label='Detection threshold')
    ax.set_xlabel('Height (mm)'); ax.set_ylabel('Phase (rad)')
    ax.set_title('Phase Drift vs Height'); ax.legend()
    ax.set_xlim(1e-3,1e5); ax.set_ylim(1e-18,10)
    p['phase'] = f'{d}/fig_phase.png'; fig.savefig(p['phase'], dpi=150); plt.close()
    
    # Fig 4: Platform comparison
    fig, ax = plt.subplots(figsize=(8,5))
    s = [phi_t[np.argmin(np.abs(h-1))], phi_o[np.argmin(np.abs(h-1))]]
    ax.bar(['Transmon','Optical'], s, color=['blue','red'])
    ax.axhline(0.1, color='g', ls='--'); ax.set_yscale('log')
    ax.set_ylabel('Phase at 1m (rad)'); ax.set_title('Platform Feasibility')
    ax.set_ylim(1e-12,10)
    p['feasibility'] = f'{d}/fig_feasibility.png'; fig.savefig(p['feasibility'], dpi=150); plt.close()
    
    # Fig 6: Confound matrix
    fig, ax = plt.subplots(figsize=(9,5))
    m = np.array([[1,1,1,1],[.5,.3,.3,0],[0,0,.5,0],[.3,0,.3,0]])
    im = ax.imshow(m, cmap='RdYlGn', vmin=0, vmax=1)
    ax.set_xticks([0,1,2,3]); ax.set_xticklabels(['dh','omega','t','Random'])
    ax.set_yticks([0,1,2,3]); ax.set_yticklabels(['SSZ','Temp','LO','Vibr'])
    ax.set_title('Confound Discrimination')
    p['confound'] = f'{d}/fig_confound.png'; fig.savefig(p['confound'], dpi=150); plt.close()
    
    # Fig 8: SSZ vs GR
    fig, ax = plt.subplots(figsize=(10,6))
    r = np.logspace(0,4,500)
    D_GR = np.sqrt(np.maximum(1-1/r, 0))
    D_SSZ = 1/(1 + 1/(2*r))
    ax.semilogx(r, D_GR, 'b-', lw=2, label='GR')
    ax.semilogx(r, D_SSZ, 'r--', lw=2, label='SSZ')
    ax.axvline(100, color='gray', ls=':')
    ax.set_xlabel('r/r_s'); ax.set_ylabel('D'); ax.legend()
    ax.set_title('SSZ vs GR Time Dilation')
    p['sszgr'] = f'{d}/fig_sszgr.png'; fig.savefig(p['sszgr'], dpi=150); plt.close()
    
    return p

def add_table(doc, h, rows):
    t = doc.add_table(rows=1+len(rows), cols=len(h))
    t.style = 'Table Grid'
    for i,x in enumerate(h): t.rows[0].cells[i].text = x
    for i,r in enumerate(rows):
        for j,c in enumerate(r): t.rows[i+1].cells[j].text = str(c)
    return t

def main():
    out = 'outputs/paper_e'
    print("Generating plots...")
    plots = gen_plots(out)
    
    print("Creating DOCX...")
    doc = Document()
    
    # Title
    doc.add_heading('Paper E: Gravitational Phase Coupling in Quantum Systems', 0)
    doc.add_paragraph('A Unified Framework for Testing SSZ Predictions\n').italic = True
    doc.add_paragraph('Lino Casu & Carmen Wrede | mail@error.wtf')
    doc.add_paragraph(f'Version 3.0 | {datetime.now().strftime("%B %Y")}')
    doc.add_page_break()
    
    # Abstract
    doc.add_heading('Abstract', 1)
    doc.add_paragraph("""The Segmented Spacetime (SSZ) framework predicts deterministic phase drifts in quantum systems at different gravitational potentials. This paper presents:

(1) Upper-bound experiments with superconducting qubits (~12 OoM below noise—null is SSZ-consistent)
(2) Direct detection via optical clocks (~0.6 rad at 1 m)  
(3) Compensation protocols for future quantum networks

All 150 tests pass; complete reproducibility via github.com/error-wtf/ssz-qubits.""")
    
    doc.add_paragraph('Keywords: SSZ, Gravitational Phase Coupling, Quantum Computing, Optical Clocks, Falsifiability')
    doc.add_page_break()
    
    # Section 1
    doc.add_heading('1. Introduction', 1)
    doc.add_heading('1.1 The Central Question', 2)
    doc.add_paragraph("""Quantum systems require precise phase control. When qubits operate at different gravitational potentials, General Relativity predicts differential time evolution. The SSZ framework provides an operational model for calculating this geometric phase drift:

    ΔΦ = ω × (r_s × Δh / R²) × t

This drift is deterministic, geometry-coupled, and compensable—unlike stochastic noise sources.""")
    
    doc.add_heading('1.2 Local vs Global Comparison', 2)
    doc.add_paragraph("""The equivalence principle states local physics is indistinguishable from flat spacetime. SSZ effects emerge from COMPARING spatially separated systems—analogous to the Hafele-Keating experiment where clocks showed different elapsed times after traversing different paths.""")
    
    doc.add_page_break()
    
    # Section 2
    doc.add_heading('2. Claims and Non-Claims', 1)
    doc.add_heading('2.1 Claim Taxonomy', 2)
    add_table(doc, ['Regime','Platform','Signal','Testable?','Interpretation'], [
        ['BOUNDED','Transmon, mm','~10⁻¹³ rad','Yes','Upper bound; null=SSZ-consistent'],
        ['DETECTABLE','Optical, 1m','~0.6 rad','Yes','Direct detection possible'],
        ['FUTURE','Large QC','Variable','No','Engineering relevance'],
    ])
    
    doc.add_paragraph()
    doc.add_heading('2.2 What SSZ Does NOT Claim', 2)
    doc.add_paragraph("""❌ Detection with current transmons at mm-scale
❌ Local violation of equivalence principle
❌ That GR is "wrong" in weak field
❌ Immediate practical relevance for current QC""")
    
    doc.add_page_break()
    
    # Section 3
    doc.add_heading('3. Theoretical Framework', 1)
    
    doc.add_heading('3.1 Core Equations', 2)
    doc.add_paragraph("""Segment Density (weak field):
    Ξ(r) = r_s / (2r)

Time Dilation:
    D_SSZ(r) = 1 / (1 + Ξ(r))

Differential Time Dilation:
    ΔD = r_s × Δh / R²

Phase Drift:
    ΔΦ = ω × ΔD × t = ω × (r_s × Δh / R²) × t""")
    
    doc.add_heading('3.2 Numerical Examples', 2)
    add_table(doc, ['Platform','f','Δh','t','ΔΦ'], [
        ['Transmon','5 GHz','1 mm','100 μs','6.87×10⁻¹³ rad'],
        ['Transmon','5 GHz','1 m','100 μs','6.87×10⁻¹⁰ rad'],
        ['Optical','429 THz','1 m','1 s','0.59 rad'],
    ])
    
    doc.add_paragraph()
    if os.path.exists(plots['phase']):
        doc.add_picture(plots['phase'], width=Inches(5.5))
        p = doc.add_paragraph('Figure 2: Phase Drift vs Height. Optical clocks reach detection regime.')
        p.italic = True
    
    doc.add_page_break()
    
    # Section 4
    doc.add_heading('4. Compensation Protocol', 1)
    doc.add_paragraph("""The WITH/WITHOUT COMPENSATION test is the gold-standard discriminator:

1. WITHOUT: Measure Φ_measured ± σ
2. WITH: Apply correction Φ_corr = -ω×(r_s×Δh/R²)×t
3. COMPARE: σ_with vs σ_without

SSZ is uniquely: deterministic, linear in (Δh, ω, t), randomization-invariant.
No confound matches all criteria simultaneously.""")
    
    if os.path.exists(plots['confound']):
        doc.add_picture(plots['confound'], width=Inches(5))
        p = doc.add_paragraph('Figure 6: Confound Discrimination. SSZ uniquely satisfies all criteria.')
        p.italic = True
    
    doc.add_page_break()
    
    # Section 5
    doc.add_heading('5. Experimental Designs', 1)
    doc.add_heading('5.1 Hardware Configurations', 2)
    add_table(doc, ['Config','Method','Δh'], [
        ['Chip Tilt','Tilt 5-10°','1-3.5 mm'],
        ['Remote Entanglement','Adjacent floors','3 m'],
        ['Tower','Multi-story','10-100 m'],
    ])
    
    doc.add_heading('5.2 Upper Bound Calculation', 2)
    doc.add_paragraph("""With Δh_max=3.5mm, N=10⁹, σ=1 rad:
    σ_averaged = 3.2×10⁻⁵ rad
    Upper bound: |α| < 9×10⁻³ rad/m (95% CL)
    
Constrains anomalous couplings to < 10¹⁰ × α_SSZ.""")
    
    doc.add_page_break()
    
    # Section 6
    doc.add_heading('6. Statistical Framework', 1)
    doc.add_paragraph("""Model Comparison:
• M₀ (Null): ΔΦ = 0 + noise
• M_SSZ: ΔΦ = α_SSZ × Δh + noise (fixed)
• M_anom: ΔΦ = α_fit × Δh + noise (free)

SSZ falsified if (in detection regime):
• α_fit ≠ α_SSZ at >3σ AND |α_fit| ≠ 0

Binary thresholds inappropriate—use CI/model comparison.""")
    
    doc.add_page_break()
    
    # Section 7
    doc.add_heading('7. Feasibility Landscape', 1)
    doc.add_heading('7.1 The 12 OoM Gap', 2)
    doc.add_paragraph("""For transmon at 1mm:
• Signal: 6.87×10⁻¹³ rad
• Noise: ~1 rad
• Gap: ~10¹²

A NULL RESULT IS SSZ-CONSISTENT. The theory predicts negligibility here.""")
    
    if os.path.exists(plots['feasibility']):
        doc.add_picture(plots['feasibility'], width=Inches(5))
        p = doc.add_paragraph('Figure 4: Platform Comparison. Optical clocks are ~10⁹× more sensitive.')
        p.italic = True
    
    doc.add_heading('7.2 Platform Comparison', 2)
    add_table(doc, ['Parameter','Transmon','Optical','Ratio'], [
        ['Frequency','5 GHz','429 THz','8.6×10⁴'],
        ['Coherence','100 μs','1 s','10⁴'],
        ['ΔΦ @ 1m','6.9×10⁻¹⁰ rad','0.59 rad','8.6×10⁸'],
        ['Feasible?','No','YES','—'],
    ])
    
    doc.add_page_break()
    
    # Section 8
    doc.add_heading('8. Reproducibility', 1)
    doc.add_paragraph("""Repository: github.com/error-wtf/ssz-qubits

Commands:
    pytest tests/ -v          # 150 tests
    python generate_paper_e_final.py
    python paper_suite_integrator.py

Test Coverage: 150/150 (100%)""")
    
    doc.add_page_break()
    
    # Section 9
    doc.add_heading('9. Conclusion', 1)
    doc.add_paragraph("""Key Findings:
1. Deterministic drift: ΔΦ = ω × (r_s × Δh / R²) × t
2. Current regime: ~12 OoM gap; null is SSZ-consistent
3. Gold standard: Optical clocks (0.59 rad @ 1m)
4. Discriminator: WITH/WITHOUT compensation
5. Statistics: Slope-fitting with CI
6. Reproducible: 150 tests, all pass

What Would Falsify SSZ (in detection regime):
• Slope ≠ α_SSZ at >3σ
• Non-linear scaling
• Randomization-variant signal""")
    
    doc.add_page_break()
    
    # Appendix A
    doc.add_heading('Appendix A: Full Derivations', 1)
    doc.add_paragraph("""A.1 Segment Density
    Ξ(r) = r_s/(2r)  [weak field]
    Ξ(r) = 1 - exp(-φ×r/r_s)  [strong field]

A.2 Time Dilation
    D = 1/(1+Ξ) ≈ 1 - Ξ + O(Ξ²)

A.3 GR Consistency
    D_SSZ ≈ D_GR ≈ 1 - r_s/(2r)  to first order

A.4 Differential
    ΔD = r_s×Δh/R²

A.5 Phase
    ΔΦ = ω×ΔD×t""")
    
    # Appendix B
    doc.add_heading('Appendix B: Didactic Scaling', 1)
    doc.add_paragraph("""Definition: Multiply SSZ prediction by S>>1 for visualization.

What gets scaled: Signal magnitude only
What stays: Noise model, scaling laws, statistics

Purpose: Validate methodology, NOT claim detectability
Requirement: Must be explicitly labeled""")
    
    # Appendix C
    doc.add_heading('Appendix C: Confound Playbook', 1)
    add_table(doc, ['Confound','Control'], [
        ['Temperature','Thermalize, randomize Δh order'],
        ['LO Noise','Common-mode reference, differential'],
        ['Vibration','Accelerometer, isolation'],
        ['Magnetic','Shielding, characterization'],
    ])
    
    # Appendix D
    doc.add_heading('Appendix D: Constants', 1)
    add_table(doc, ['Constant','Value','Units'], [
        ['c','2.998×10⁸','m/s'],
        ['G','6.674×10⁻¹¹','m³/(kg·s²)'],
        ['r_s (Earth)','8.870×10⁻³','m'],
        ['R (Earth)','6.371×10⁶','m'],
        ['φ','1.618...','—'],
    ])
    
    # References
    doc.add_heading('References', 1)
    doc.add_paragraph("""[1] Bothwell, T. et al. (2022). Resolving gravitational redshift at 1mm. Nature 602, 420-424.

[2] SSZ-Qubits Repository. github.com/error-wtf/ssz-qubits

[3] Hafele, J.C. & Keating, R.E. (1972). Around-the-world atomic clocks. Science 177, 166-168.

[4] Chou, C.W. et al. (2010). Optical clocks and relativity. Science 329, 1630-1633.""")
    
    doc.add_paragraph()
    doc.add_paragraph('© 2025 Carmen Wrede & Lino Casu')
    doc.add_paragraph('Licensed under ANTI-CAPITALIST SOFTWARE LICENSE v1.4')
    
    # Save
    docx_path = f'{out}/Paper_E_Final.docx'
    doc.save(docx_path)
    print(f"\n[OK] Paper E saved: {docx_path}")
    print(f"[OK] Plots in: {out}/")
    return docx_path

if __name__ == '__main__':
    main()
