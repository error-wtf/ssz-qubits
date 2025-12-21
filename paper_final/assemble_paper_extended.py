#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
SSZ Unified Paper Assembler - EXTENDED VERSION
Generates final DOCX with all extended explanatory content.

Usage:
    python assemble_paper_extended.py

Output:
    paper_final/output/SSZ_Unified_Paper_EXTENDED.docx

© 2025 Carmen Wrede & Lino Casu
Licensed under ANTI-CAPITALIST SOFTWARE LICENSE v1.4
"""

import os
import sys
import subprocess
from pathlib import Path
from datetime import datetime

# Ensure UTF-8 for Windows
os.environ['PYTHONIOENCODING'] = 'utf-8:replace'

# Check dependencies
try:
    from docx import Document
    from docx.shared import Inches, Pt, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.style import WD_STYLE_TYPE
except ImportError:
    print("Installing python-docx...")
    subprocess.run([sys.executable, "-m", "pip", "install", "python-docx"], check=True)
    from docx import Document
    from docx.shared import Inches, Pt, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH

try:
    import matplotlib
    matplotlib.use('Agg')
    import matplotlib.pyplot as plt
    import numpy as np
except ImportError:
    print("Installing matplotlib...")
    subprocess.run([sys.executable, "-m", "pip", "install", "matplotlib", "numpy"], check=True)
    import matplotlib
    matplotlib.use('Agg')
    import matplotlib.pyplot as plt
    import numpy as np

# Paths
BASE_DIR = Path(__file__).parent
SECTIONS_DIR = BASE_DIR / "sections"
SECTIONS_EXTENDED_DIR = BASE_DIR / "sections_extended"
TABLES_DIR = BASE_DIR / "tables"
FIGURES_DIR = BASE_DIR / "figures"
APPENDICES_DIR = BASE_DIR / "appendices"
APPENDICES_EXTENDED_DIR = BASE_DIR / "appendices_extended"
OUTPUT_DIR = BASE_DIR / "output"
OUTPUT_DIR.mkdir(exist_ok=True)
(OUTPUT_DIR / "figures").mkdir(exist_ok=True)


# Mapping: which extended files supplement which original sections
SECTION_EXTENSIONS = {
    "01_introduction": "01_claim_boundaries_extended.md",
    "01b_history": "01b_history_extended.md",
    "02b_strong_field": "02b_strong_field_extended.md",
    "03_qubit_physics": "03_qubit_physics_extended.md",
    "03_control": "03_control_extended.md",
    "04_experiments": "04_experiments_extended.md",
    "05_entanglement": "05_entanglement_extended.md",
    "06_engineering": "06_engineering_extended.md",
    "07_feasibility": "07_feasibility_extended.md",
    "08_conclusion": "08_conclusion_extended.md",
    "10_roadmap": "10_roadmap_extended.md",
}

APPENDIX_EXTENSIONS = {
    "A_derivation": "A_derivation_extended.md",
    "C_confounds": "C_confounds_extended.md",
}


def generate_all_figures():
    """Generate all figures by running figure scripts."""
    print("Generating figures...")
    
    figure_scripts = list(FIGURES_DIR.glob("F*.py"))
    for script in figure_scripts:
        print(f"  Running {script.name}...")
        try:
            result = subprocess.run(
                [sys.executable, str(script)],
                cwd=str(OUTPUT_DIR / "figures"),
                capture_output=True,
                text=True,
                encoding='utf-8',
                errors='replace',
                timeout=60
            )
            if result.returncode != 0:
                print(f"    Warning: {script.name} failed: {result.stderr[:200]}")
        except Exception as e:
            print(f"    Error running {script.name}: {e}")
    
    print(f"  Figures saved to {OUTPUT_DIR / 'figures'}/")


def read_section(name, use_extended=True):
    """Read a section markdown file, optionally with extended content."""
    # Read original section
    path = SECTIONS_DIR / f"{name}.md"
    content = ""
    if path.exists():
        with open(path, 'r', encoding='utf-8') as f:
            content = f.read()
    
    # Append extended content if available
    if use_extended and name in SECTION_EXTENSIONS:
        ext_path = SECTIONS_EXTENDED_DIR / SECTION_EXTENSIONS[name]
        if ext_path.exists():
            with open(ext_path, 'r', encoding='utf-8') as f:
                extended = f.read()
            # Add separator and extended content
            content += "\n\n---\n\n## Erweiterte Ausführungen\n\n"
            content += extended
            print(f"    + Extended content from {SECTION_EXTENSIONS[name]}")
    
    return content


def read_table(name):
    """Read a table markdown file."""
    path = TABLES_DIR / f"{name}.md"
    if path.exists():
        with open(path, 'r', encoding='utf-8') as f:
            return f.read()
    return ""


def read_appendix(name, use_extended=True):
    """Read an appendix markdown file, optionally with extended content."""
    # Read original appendix
    path = APPENDICES_DIR / f"{name}.md"
    content = ""
    if path.exists():
        with open(path, 'r', encoding='utf-8') as f:
            content = f.read()
    
    # Append extended content if available
    if use_extended and name in APPENDIX_EXTENSIONS:
        ext_path = APPENDICES_EXTENDED_DIR / APPENDIX_EXTENSIONS[name]
        if ext_path.exists():
            with open(ext_path, 'r', encoding='utf-8') as f:
                extended = f.read()
            content += "\n\n---\n\n## Erweiterte Erläuterungen\n\n"
            content += extended
            print(f"    + Extended content from {APPENDIX_EXTENSIONS[name]}")
    
    return content


def add_heading(doc, text, level=1):
    """Add a heading to the document."""
    doc.add_heading(text, level=level)


def add_paragraph(doc, text, bold=False, italic=False):
    """Add a paragraph to the document."""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    return p


def add_markdown_content(doc, md_text):
    """Parse markdown and add to document (simplified)."""
    lines = md_text.split('\n')
    in_code_block = False
    code_content = []
    
    for line in lines:
        # Skip empty lines at start
        if not line.strip() and not doc.paragraphs:
            continue
            
        # Code blocks
        if line.strip().startswith('```'):
            if in_code_block:
                # End code block
                code_text = '\n'.join(code_content)
                p = doc.add_paragraph()
                p.style = 'No Spacing'
                run = p.add_run(code_text)
                run.font.name = 'Courier New'
                run.font.size = Pt(9)
                code_content = []
                in_code_block = False
            else:
                in_code_block = True
            continue
        
        if in_code_block:
            code_content.append(line)
            continue
        
        # Headings
        if line.startswith('# '):
            add_heading(doc, line[2:].strip(), level=1)
        elif line.startswith('## '):
            add_heading(doc, line[3:].strip(), level=2)
        elif line.startswith('### '):
            add_heading(doc, line[4:].strip(), level=3)
        elif line.startswith('#### '):
            add_heading(doc, line[5:].strip(), level=4)
        # Horizontal rules
        elif line.strip() == '---':
            doc.add_paragraph('─' * 50)
        # Tables (simplified - just add as text)
        elif line.startswith('|'):
            p = doc.add_paragraph(line)
            p.style = 'No Spacing'
            for run in p.runs:
                run.font.name = 'Courier New'
                run.font.size = Pt(9)
        # Lists
        elif line.strip().startswith('- ') or line.strip().startswith('* '):
            text = line.strip()[2:]
            p = doc.add_paragraph(text, style='List Bullet')
        elif line.strip() and line.strip()[0].isdigit() and '. ' in line:
            text = line.strip().split('. ', 1)[1] if '. ' in line else line.strip()
            p = doc.add_paragraph(text, style='List Number')
        # Regular paragraph
        elif line.strip():
            text = line.strip()
            add_paragraph(doc, text)
        # Empty line
        else:
            doc.add_paragraph()


def add_figure(doc, figure_name, caption=""):
    """Add a figure to the document."""
    # First try source figures directory (has enhanced versions)
    fig_path = FIGURES_DIR / f"{figure_name}.png"
    if not fig_path.exists():
        # Fallback to output directory
        fig_path = OUTPUT_DIR / "figures" / f"{figure_name}.png"
    
    if fig_path.exists():
        doc.add_picture(str(fig_path), width=Inches(6))
        if caption:
            p = doc.add_paragraph(caption)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.runs[0].italic = True
    else:
        add_paragraph(doc, f"[Figure {figure_name} not generated]", italic=True)


def create_title_page(doc):
    """Create the title page."""
    # Title
    title = doc.add_heading('Segmented Spacetime:', level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    subtitle = doc.add_heading('Gravitational Phase Coupling in Quantum Systems', level=1)
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Extended edition marker
    doc.add_paragraph()
    extended_marker = doc.add_paragraph('EXTENDED EDITION')
    extended_marker.alignment = WD_ALIGN_PARAGRAPH.CENTER
    extended_marker.runs[0].bold = True
    extended_marker.runs[0].font.size = Pt(14)
    
    doc.add_paragraph()
    
    # Authors
    authors = doc.add_paragraph('Carmen Wrede & Lino Casu')
    authors.alignment = WD_ALIGN_PARAGRAPH.CENTER
    authors.runs[0].bold = True
    
    # Affiliation placeholder
    affil = doc.add_paragraph('Independent Researchers')
    affil.alignment = WD_ALIGN_PARAGRAPH.CENTER
    affil.runs[0].italic = True
    
    doc.add_paragraph()
    
    # Date
    date_p = doc.add_paragraph(datetime.now().strftime('%B %Y'))
    date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()
    
    # Version info
    version = doc.add_paragraph('Extended version with detailed explanations')
    version.alignment = WD_ALIGN_PARAGRAPH.CENTER
    version.runs[0].italic = True
    
    doc.add_paragraph()
    
    # Repository
    repo = doc.add_paragraph('Repository: github.com/error-wtf/ssz-qubits')
    repo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # License
    lic = doc.add_paragraph('Licensed under ANTI-CAPITALIST SOFTWARE LICENSE v1.4')
    lic.alignment = WD_ALIGN_PARAGRAPH.CENTER
    lic.runs[0].font.size = Pt(10)
    
    doc.add_page_break()


def create_toc_placeholder(doc):
    """Add a table of contents placeholder."""
    add_heading(doc, 'Table of Contents', level=1)
    
    toc_items = [
        "1. Introduction and Scope (+ Extended Claim Boundaries)",
        "2. Theory: Deriving the SSZ Phase Drift",
        "   2.5 Strong Field Extension (+ Why φ?)",
        "3. Control and Compensation Protocol (+ Extended)",
        "4. Experimental Designs and Feasibility (+ Extended)",
        "5. Entanglement and Phase Preservation (+ Lindblad)",
        "6. Engineering Implications (+ Extended)",
        "7. Feasibility Landscape (+ Roadmap Extended)",
        "8. Conclusion (+ Extended 5 Points)",
        "Appendix A: Full Mathematical Derivation (+ Extended)",
        "Appendix B: Didactic Scaling Definition",
        "Appendix C: Confound Playbook (+ Extended)",
        "Appendix D: Physical Constants",
        "Appendix E-H: Additional Appendices",
        "References",
    ]
    
    for item in toc_items:
        p = doc.add_paragraph(item)
        p.paragraph_format.left_indent = Inches(0.5)
    
    doc.add_page_break()


def assemble_document():
    """Assemble the full extended document."""
    print("Assembling EXTENDED document...")
    
    doc = Document()
    
    # Title page
    create_title_page(doc)
    
    # TOC
    create_toc_placeholder(doc)
    
    # Abstract
    print("  Adding Abstract...")
    add_markdown_content(doc, read_section("00_abstract", use_extended=False))
    doc.add_page_break()
    
    # Main sections with extensions
    sections = [
        ("01_introduction", "F11_claim_boundaries", True),
        ("01b_history", "F9_validation", True),
        ("02_theory", "F1_phase_vs_height", False),
        ("02b_strong_field", "F7_strong_field", True),
        ("03_qubit_physics", None, True),
        ("03_control", "F6_compensation", True),
        ("04_experiments", "F2_platform_comparison", True),
        ("05_entanglement", None, True),
        ("06_engineering", "F5_zone_width", True),
        ("07_feasibility", "F3_confound_matrix", True),
        ("08_conclusion", None, True),
        ("10_roadmap", "F8_timeline", True),
        ("12_reproducibility", "F10_network", False),
    ]
    
    for section_name, figure_name, has_extension in sections:
        print(f"  Adding {section_name}..." + (" [EXTENDED]" if has_extension else ""))
        add_markdown_content(doc, read_section(section_name, use_extended=has_extension))
        
        if figure_name:
            doc.add_paragraph()
            add_figure(doc, figure_name, f"Figure: {figure_name.replace('_', ' ').title()}")
        
        doc.add_page_break()
    
    # Appendices with extensions
    appendices = [
        ("A_derivation", "Appendix A: Full Mathematical Derivation", True),
        ("B_didactic", "Appendix B: Didactic Scaling Definition", False),
        ("C_confounds", "Appendix C: Confound Playbook", True),
        ("D_constants", "Appendix D: Physical Constants", False),
        ("E_transition", "Appendix E: Weak-Strong Field Transition", False),
        ("F_platforms", "Appendix F: Platform Technical Specifications", False),
        ("G_statistics", "Appendix G: Statistical Methods", False),
        ("H_code", "Appendix H: Code Listings", False),
    ]
    
    for app_name, app_title, has_extension in appendices:
        print(f"  Adding {app_name}..." + (" [EXTENDED]" if has_extension else ""))
        add_markdown_content(doc, read_appendix(app_name, use_extended=has_extension))
        doc.add_page_break()
    
    # References
    print("  Adding References...")
    ref_path = BASE_DIR / "references.md"
    if ref_path.exists():
        with open(ref_path, 'r', encoding='utf-8') as f:
            add_markdown_content(doc, f.read())
    
    # Save
    output_path = OUTPUT_DIR / "SSZ_Unified_Paper_EXTENDED.docx"
    doc.save(str(output_path))
    print(f"\n[OK] Document saved: {output_path}")
    print(f"[OK] Size: {output_path.stat().st_size / 1024:.1f} KB")
    
    return output_path


def main():
    """Main entry point."""
    print("=" * 60)
    print("SSZ UNIFIED PAPER ASSEMBLER - EXTENDED VERSION")
    print("=" * 60)
    print()
    
    # Show extended content summary
    print("Extended content to be included:")
    for section, ext_file in SECTION_EXTENSIONS.items():
        ext_path = SECTIONS_EXTENDED_DIR / ext_file
        if ext_path.exists():
            size = ext_path.stat().st_size / 1024
            print(f"  + {section}: {ext_file} ({size:.1f} KB)")
    for appendix, ext_file in APPENDIX_EXTENSIONS.items():
        ext_path = APPENDICES_EXTENDED_DIR / ext_file
        if ext_path.exists():
            size = ext_path.stat().st_size / 1024
            print(f"  + {appendix}: {ext_file} ({size:.1f} KB)")
    print()
    
    # Skip figure generation - use pre-generated enhanced figures from figures/ directory
    print("Using pre-generated figures from figures/ directory...")
    print()
    
    # Assemble document
    output = assemble_document()
    
    print()
    print("=" * 60)
    print("EXTENDED ASSEMBLY COMPLETE")
    print("=" * 60)
    print(f"Output: {output}")
    print()
    print("Components used:")
    print(f"  - {len(list(SECTIONS_DIR.glob('*.md')))} sections")
    print(f"  - {len(list(SECTIONS_EXTENDED_DIR.glob('*.md')))} extended sections")
    print(f"  - {len(list(TABLES_DIR.glob('*.md')))} tables")
    print(f"  - {len(list(FIGURES_DIR.glob('*.py')))} figures")
    print(f"  - {len(list(APPENDICES_DIR.glob('*.md')))} appendices")
    print(f"  - {len(list(APPENDICES_EXTENDED_DIR.glob('*.md')))} extended appendices")
    print()
    
    return 0


if __name__ == '__main__':
    sys.exit(main())
