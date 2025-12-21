#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Generate Paper C v1.1 DOCX Document

Upper-Bound / Feasibility-Corrected Version

(c) 2025 Carmen Wrede, Lino Casu
"""

import os
import sys
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))
OUTPUT_DIR = os.path.join(SCRIPT_DIR, 'outputs')
PAPERS_DIR = r'E:\clone\SSZ_QUBIT_PAPERS'

def set_cell_shading(cell, color):
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color)
    cell._tc.get_or_add_tcPr().append(shading)

def create_paper_c_v11():
    doc = Document()
    
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(11)
    
    # Title
    title = doc.add_paragraph()
    run = title.add_run('Experimental Framework for Testing Gravitational Phase Coupling in Quantum Systems')
    run.bold = True
    run.font.size = Pt(16)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    subtitle = doc.add_paragraph()
    run = subtitle.add_run('A Protocol for Upper-Bound Constraints and Platform Comparison')
    run.italic = True
    run.font.size = Pt(12)
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    version = doc.add_paragraph()
    version.add_run('Paper C v1.1 (Revised with Feasibility Analysis)').bold = True
    version.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()
    
    authors = doc.add_paragraph()
    authors.add_run('Lino Casu, Carmen Wrede').bold = True
    authors.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph('Independent Researchers').alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph('Contact: mail@error.wtf').alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph('December 2025').alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()
    
    # Abstract
    doc.add_heading('Abstract', level=1)
    
    abstract = """We present an experimental framework for testing gravitational phase coupling in quantum systems, as predicted by the Segmented Spacetime (SSZ) model. Building on Papers A and B, we perform a rigorous order-of-magnitude feasibility analysis revealing that the predicted SSZ effect at laboratory-scale height differences (Dh ~ mm) is approximately 12 orders of magnitude below the noise floor of current superconducting qubit technology. We therefore reframe this work as: (1) an upper-bound experiment that can constrain anomalous phase couplings in solid-state qubits, (2) a platform comparison identifying optical clocks as the appropriate gold-standard test system, and (3) a statistical falsification framework using slope-fitting rather than binary thresholds. We provide concrete hardware implementations for height-difference generation (chip tilt, remote entanglement, 3D chiplet stacks) and a confound discrimination strategy based on scaling signatures rather than absolute exclusions."""
    
    p = doc.add_paragraph(abstract)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run('Keywords: ').bold = True
    p.add_run('Gravitational Phase Coupling, Quantum Systems, Upper Bound, Feasibility Analysis, Optical Clocks, Falsifiability')
    
    doc.add_page_break()
    
    # 1. Introduction
    doc.add_heading('1. Introduction', level=1)
    
    doc.add_heading('1.1 Context from Papers A and B', level=2)
    p = doc.add_paragraph('In Papers A and B, we derived the SSZ prediction for gravitational phase drift:')
    
    eq = doc.add_paragraph()
    eq.add_run('DF(t) = w x DD').italic = True
    eq.add_run('SSZ').font.subscript = True
    eq.add_run('(Dh) x t')
    eq.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    p = doc.add_paragraph('where DD')
    p.add_run('SSZ').font.subscript = True
    p.add_run('(Dh) = r')
    p.add_run('s').font.subscript = True
    p.add_run(' x Dh / R')
    p.add_run('2').font.superscript = True
    p.add_run(' for small height differences.')
    
    doc.add_heading('1.2 The Critical Question', level=2)
    p = doc.add_paragraph()
    p.add_run('Can this effect be detected with current technology?').bold = True
    
    p = doc.add_paragraph('This paper provides the honest answer: ')
    p.add_run('No, not at laboratory scales with superconducting qubits.').bold = True
    p.add_run(' However, this negative result is scientifically valuable when framed correctly.')
    
    doc.add_heading('1.3 Revised Goals', level=2)
    doc.add_paragraph('Quantify the feasibility gap between predicted signal and noise floor', style='List Number')
    doc.add_paragraph('Identify appropriate experimental platforms where detection is possible', style='List Number')
    doc.add_paragraph('Design an upper-bound experiment that provides value regardless of outcome', style='List Number')
    doc.add_paragraph('Establish a statistical framework for falsification claims', style='List Number')
    
    # 2. Feasibility Analysis
    doc.add_heading('2. Feasibility Analysis', level=1)
    
    doc.add_heading('2.1 Signal Size', level=2)
    p = doc.add_paragraph('For a 5 GHz qubit with Ramsey time T = 100 us:')
    
    table = doc.add_table(rows=5, cols=3)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    headers = ['Dh', 'DD_SSZ', 'DF (100 us)']
    for i, h in enumerate(headers):
        table.rows[0].cells[i].text = h
        table.rows[0].cells[i].paragraphs[0].runs[0].bold = True
        set_cell_shading(table.rows[0].cells[i], 'D9E2F3')
    
    data = [
        ('1 mm', '1.09x10^-19', '3.43x10^-13 rad'),
        ('1 m', '1.09x10^-16', '3.43x10^-10 rad'),
        ('10 m', '1.09x10^-15', '3.43x10^-9 rad'),
        ('100 m', '1.09x10^-14', '3.43x10^-8 rad'),
    ]
    for i, row in enumerate(data):
        for j, val in enumerate(row):
            table.rows[i+1].cells[j].text = val
    
    doc.add_paragraph()
    
    doc.add_heading('2.2 Noise Floor', level=2)
    p = doc.add_paragraph('State-of-the-art superconducting qubit phase measurement:')
    
    table2 = doc.add_table(rows=5, cols=2)
    table2.style = 'Table Grid'
    table2.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    table2.rows[0].cells[0].text = 'Source'
    table2.rows[0].cells[1].text = 'Magnitude (single shot)'
    for cell in table2.rows[0].cells:
        cell.paragraphs[0].runs[0].bold = True
        set_cell_shading(cell, 'D9E2F3')
    
    noise_data = [
        ('Quantum projection noise', '~1 rad'),
        ('LO phase noise (100 us)', '~10^-3 rad'),
        ('Temperature drift (1 mK)', '~0.6 rad'),
        ('Combined', '~1 rad'),
    ]
    for i, (src, mag) in enumerate(noise_data):
        table2.rows[i+1].cells[0].text = src
        table2.rows[i+1].cells[1].text = mag
    
    doc.add_paragraph()
    
    doc.add_heading('2.3 Averaging Requirements', level=2)
    
    p = doc.add_paragraph('For SNR = 3, the required number of shots:')
    
    table3 = doc.add_table(rows=4, cols=5)
    table3.style = 'Table Grid'
    table3.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    headers3 = ['Dh', 'Signal', 'N required', 'Time @ 10 kHz', 'Feasible?']
    for i, h in enumerate(headers3):
        table3.rows[0].cells[i].text = h
        table3.rows[0].cells[i].paragraphs[0].runs[0].bold = True
        set_cell_shading(table3.rows[0].cells[i], 'D9E2F3')
    
    avg_data = [
        ('1 mm', '3.4x10^-13 rad', '7.6x10^25', '2.4x10^14 years', 'No'),
        ('1 m', '3.4x10^-10 rad', '7.6x10^19', '2.4x10^8 years', 'No'),
        ('10 m', '3.4x10^-9 rad', '7.6x10^17', '2.4x10^6 years', 'No'),
    ]
    for i, row in enumerate(avg_data):
        for j, val in enumerate(row):
            table3.rows[i+1].cells[j].text = val
    
    doc.add_paragraph()
    
    doc.add_heading('2.4 Conclusion', level=2)
    p = doc.add_paragraph()
    p.add_run('The SSZ effect at GR-predicted levels is ~12 orders of magnitude below detectability with current superconducting qubit technology.').bold = True
    
    p = doc.add_paragraph('This is not a failure of the theory--it is the expected regime where gravitational effects are negligible for solid-state systems on Earth.')
    
    doc.add_page_break()
    
    # 3. Alternative Platforms
    doc.add_heading('3. Alternative Platforms', level=1)
    
    doc.add_heading('3.1 Optical Atomic Clocks', level=2)
    
    table4 = doc.add_table(rows=6, cols=3)
    table4.style = 'Table Grid'
    table4.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    headers4 = ['Parameter', 'Transmon Qubit', 'Optical Clock']
    for i, h in enumerate(headers4):
        table4.rows[0].cells[i].text = h
        table4.rows[0].cells[i].paragraphs[0].runs[0].bold = True
        set_cell_shading(table4.rows[0].cells[i], 'D9E2F3')
    
    platform_data = [
        ('Frequency', '5 GHz', '429 THz'),
        ('Coherence', '100 us', '1 s'),
        ('DF @ Dh=1m', '3.4x10^-10 rad', '1.3x10^-6 rad'),
        ('N for SNR=3', '7.6x10^19', '~10^9'),
        ('Feasible?', 'No', 'Yes'),
    ]
    for i, row in enumerate(platform_data):
        for j, val in enumerate(row):
            table4.rows[i+1].cells[j].text = val
    
    doc.add_paragraph()
    
    p = doc.add_paragraph()
    p.add_run('Optical clock experiments have already demonstrated gravitational redshift at the ~1 cm level').bold = True
    p.add_run(' (Bothwell et al., Nature 2022). This is the appropriate platform for testing gravitational phase coupling.')
    
    doc.add_heading('3.2 Recommendation', level=2)
    p = doc.add_paragraph()
    p.add_run('For testing SSZ predictions quantitatively, optical atomic clocks are the gold-standard platform.').bold = True
    p.add_run(' Superconducting qubits can provide upper bounds on anomalous couplings but cannot detect GR-level effects.')
    
    # 4. Upper-Bound Experiment
    doc.add_heading('4. Upper-Bound Experiment Design', level=1)
    
    doc.add_heading('4.1 Scientific Value', level=2)
    doc.add_paragraph('Constraining anomalous couplings: If any beyond-GR phase coupling exists, it must be smaller than our upper bound', style='List Bullet')
    doc.add_paragraph('Validating null predictions: SSZ predicts negligible effect at mm-scale--confirming this is a positive result', style='List Bullet')
    doc.add_paragraph('Establishing methodology: First systematic study of gravitational phase coupling in solid-state qubits', style='List Bullet')
    
    doc.add_heading('4.2 Hardware Configurations', level=2)
    
    p = doc.add_paragraph()
    p.add_run('Configuration A: Chip Tilt').bold = True
    
    table5 = doc.add_table(rows=4, cols=2)
    table5.style = 'Table Grid'
    table5.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    table5.rows[0].cells[0].text = 'Tilt Angle'
    table5.rows[0].cells[1].text = 'Dh across 20 mm chip'
    for cell in table5.rows[0].cells:
        cell.paragraphs[0].runs[0].bold = True
        set_cell_shading(cell, 'D9E2F3')
    
    tilt_data = [('1 deg', '0.35 mm'), ('5 deg', '1.74 mm'), ('10 deg', '3.47 mm')]
    for i, (angle, dh) in enumerate(tilt_data):
        table5.rows[i+1].cells[0].text = angle
        table5.rows[i+1].cells[1].text = dh
    
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run('Implementation: ').italic = True
    p.add_run('Precision goniometer stage under dilution refrigerator sample mount.')
    
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run('Configuration B: Remote Entanglement').bold = True
    
    p = doc.add_paragraph('Two qubits in separate dilution refrigerators at different heights, connected via microwave link or fiber-optical transduction.')
    
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run('Configuration C: 3D Chiplet Stack').bold = True
    
    p = doc.add_paragraph('Vertically stacked quantum processors with through-silicon vias. Emerging technology pursued by IBM and Google.')
    
    doc.add_page_break()
    
    # 5. Statistical Framework
    doc.add_heading('5. Statistical Falsification Framework', level=1)
    
    doc.add_heading('5.1 Replacing Binary Thresholds', level=2)
    p = doc.add_paragraph('The v1.0 falsification thresholds ("<50% -> falsified") are replaced with a proper statistical framework based on slope-fitting and model comparison.')
    
    doc.add_heading('5.2 Model Comparison', level=2)
    
    p = doc.add_paragraph()
    p.add_run('M').bold = True
    p.add_run('0').font.subscript = True
    p.add_run(' (Null): ').bold = True
    p.add_run('DF = 0 + noise')
    
    p = doc.add_paragraph()
    p.add_run('M').bold = True
    p.add_run('SSZ').font.subscript = True
    p.add_run(' (SSZ prediction): ').bold = True
    p.add_run('DF = a_SSZ x Dh + noise')
    
    p = doc.add_paragraph()
    p.add_run('M').bold = True
    p.add_run('anom').font.subscript = True
    p.add_run(' (Anomalous): ').bold = True
    p.add_run('DF = a_fit x Dh + noise (free parameter)')
    
    doc.add_heading('5.3 Falsification Criteria', level=2)
    
    p = doc.add_paragraph()
    p.add_run('SSZ falsified if:').bold = True
    doc.add_paragraph('Measured slope a_fit is inconsistent with a_SSZ at >3s', style='List Bullet')
    doc.add_paragraph('AND |a_fit| significantly different from zero', style='List Bullet')
    
    p = doc.add_paragraph()
    p.add_run('SSZ supported if:').bold = True
    doc.add_paragraph('Measured slope consistent with a_SSZ within uncertainty', style='List Bullet')
    doc.add_paragraph('OR null result consistent with a_SSZ ~ 0 (at mm-scale, this is the prediction!)', style='List Bullet')
    
    doc.add_heading('5.4 Upper Bound Statement', level=2)
    p = doc.add_paragraph('If no signal is detected:')
    
    eq = doc.add_paragraph('|a_anomalous| < s_slope / Dh_max (95% CL)')
    eq.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    p = doc.add_paragraph('This constrains any gravitational phase coupling to be smaller than the measurement uncertainty.')
    
    # 6. Confound Discrimination
    doc.add_heading('6. Confound Discrimination (Revised)', level=1)
    
    doc.add_heading('6.1 Principle: Signatures, Not Exclusions', level=2)
    p = doc.add_paragraph('Instead of claiming confounds "cannot" produce certain effects, we identify ')
    p.add_run('distinct scaling signatures:').bold = True
    
    table6 = doc.add_table(rows=6, cols=5)
    table6.style = 'Table Grid'
    table6.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    headers6 = ['Source', 'Dh scaling', 'w scaling', 't scaling', 'Randomization']
    for i, h in enumerate(headers6):
        table6.rows[0].cells[i].text = h
        table6.rows[0].cells[i].paragraphs[0].runs[0].bold = True
        set_cell_shading(table6.rows[0].cells[i], 'D9E2F3')
    
    confound_data = [
        ('SSZ', 'Linear', 'Linear', 'Linear', 'Invariant'),
        ('Temperature', 'Non-linear', 'Weak', 'Non-linear', 'Varies'),
        ('LO noise', 'None', 'None', 'sqrt(t)', 'Varies'),
        ('Vibration', 'Correlated', 'None', 'AC', 'Varies'),
        ('EM crosstalk', 'Position-dep.', 'Weak', 'None', 'Varies'),
    ]
    for i, row in enumerate(confound_data):
        for j, val in enumerate(row):
            table6.rows[i+1].cells[j].text = val
    
    doc.add_paragraph()
    
    doc.add_heading('6.2 The Differential Test', level=2)
    p = doc.add_paragraph('The strongest discriminator remains ')
    p.add_run('compensation:').bold = True
    
    doc.add_paragraph('Measure DF without SSZ correction', style='List Number')
    doc.add_paragraph('Apply predicted SSZ correction', style='List Number')
    doc.add_paragraph('Measure DF with correction', style='List Number')
    
    p = doc.add_paragraph()
    p.add_run('Interpretation:').bold = True
    doc.add_paragraph('If correction reduces variance: supports geometry-linked coupling', style='List Bullet')
    doc.add_paragraph('If correction has no effect: no detectable coupling', style='List Bullet')
    doc.add_paragraph('If correction increases variance: model is wrong', style='List Bullet')
    
    doc.add_page_break()
    
    # 7. Consistency
    doc.add_heading('7. Consistency with Papers A and B', level=1)
    
    doc.add_heading('7.1 Apparent Contradiction', level=2)
    p = doc.add_paragraph('Papers A/B suggest SSZ effects are relevant for quantum computing. Paper C shows they are undetectable. How to reconcile?')
    
    doc.add_heading('7.2 Resolution', level=2)
    
    p = doc.add_paragraph()
    p.add_run('Papers A/B: ').bold = True
    p.add_run('Describe the regime where SSZ becomes relevant--as QEC improves and coherence times extend, the cumulative effect grows. The papers identify when SSZ corrections would be needed (future systems).')
    
    p = doc.add_paragraph()
    p.add_run('Paper C: ').bold = True
    p.add_run('Tests current systems where SSZ effects are negligible. This is not a contradiction--it is the expected result in the present regime.')
    
    doc.add_heading('7.3 The Scaling Argument', level=2)
    p = doc.add_paragraph('Relevance scales with coherence time, qubit frequency, gate count, and height difference. For current systems (T2 ~ 100 us, Dh ~ mm) the effect is negligible. For future systems (T2 ~ 1 s, Dh ~ m) it may become relevant.')
    
    # 8. Conclusion
    doc.add_heading('8. Conclusion', level=1)
    
    p = doc.add_paragraph('We have presented a revised experimental framework:')
    
    doc.add_paragraph('Feasibility: The predicted SSZ effect is ~12 orders of magnitude below current superconducting qubit sensitivity at mm-scale Dh', style='List Number')
    doc.add_paragraph('Reframing: This paper provides an upper-bound protocol rather than a detection experiment', style='List Number')
    doc.add_paragraph('Platform: Optical atomic clocks are identified as the appropriate gold-standard platform', style='List Number')
    doc.add_paragraph('Statistics: Falsification is based on slope-fitting and model comparison, not binary thresholds', style='List Number')
    doc.add_paragraph('Value: Even null results constrain anomalous phase couplings', style='List Number')
    
    p = doc.add_paragraph()
    p.add_run('The SSZ framework makes testable predictions. This paper honestly assesses where those tests are feasible and how they should be conducted.').italic = True
    
    # References
    doc.add_heading('References', level=1)
    doc.add_paragraph('[1] Casu, L. & Wrede, C. (2025). Paper A: Geometric Qubit Optimization via Segmented Spacetime.')
    doc.add_paragraph('[2] Casu, L. & Wrede, C. (2025). Paper B: Phase Coherence and Entanglement Preservation.')
    doc.add_paragraph('[3] Bothwell, T. et al. (2022). Resolving the gravitational redshift across a millimetre-scale atomic sample. Nature 602, 420-424.')
    doc.add_paragraph('[4] SSZ-Qubits Repository: https://github.com/error-wtf/ssz-qubits')
    
    # Footer
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run('(c) 2025 Carmen Wrede & Lino Casu').italic = True
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    p = doc.add_paragraph()
    p.add_run('Licensed under the ANTI-CAPITALIST SOFTWARE LICENSE v1.4').italic = True
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Save
    os.makedirs(OUTPUT_DIR, exist_ok=True)
    os.makedirs(PAPERS_DIR, exist_ok=True)
    
    path1 = os.path.join(PAPERS_DIR, 'SSZ_Paper_C_v1.1_Upper_Bound.docx')
    path2 = os.path.join(OUTPUT_DIR, 'SSZ_Paper_C_v1.1_Upper_Bound.docx')
    
    doc.save(path1)
    doc.save(path2)
    
    print(f"Saved: {path1}")
    print(f"Saved: {path2}")
    
    return path1

if __name__ == "__main__":
    create_paper_c_v11()
