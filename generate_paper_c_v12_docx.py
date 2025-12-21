#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Generate Paper C v1.2 DOCX - Bulletproof Edition

Patches applied:
1. Optical clock fix: 0.29 rad (not 10^-6)
2. Noise: "representative" not "state-of-the-art"
3. Chip tilt: Δh = L×sin(θ) explanation
4. Upper bound: concrete example with numbers
5. Abstract: "null is positive result" sentence

(c) 2025 Carmen Wrede, Lino Casu
"""

import os
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))
OUTPUT_DIR = os.path.join(SCRIPT_DIR, 'outputs')
PAPERS_DIR = r'E:\clone\SSZ_QUBIT_PAPERS'

def set_cell_shading(cell, color):
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color)
    cell._tc.get_or_add_tcPr().append(shading)

def create_paper_c_v12():
    doc = Document()
    
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(11)
    
    # Title
    title = doc.add_paragraph()
    run = title.add_run('Experimental Framework for Testing Gravitational Phase Coupling in Quantum Systems')
    run.bold = True
    run.font.size = Pt(16)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    subtitle = doc.add_paragraph()
    run = subtitle.add_run('A Protocol for Upper-Bound Constraints and Platform Comparison')
    run.italic = True
    run.font.size = Pt(12)
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    version = doc.add_paragraph()
    version.add_run('Paper C v1.2 (Bulletproof Edition)').bold = True
    version.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()
    
    authors = doc.add_paragraph()
    authors.add_run('Lino Casu, Carmen Wrede').bold = True
    authors.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph('Independent Researchers').alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph('Contact: mail@error.wtf').alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph('December 2025').alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()
    
    # Abstract - PATCH 5: "null is positive result"
    doc.add_heading('Abstract', level=1)
    
    abstract = """We present an experimental framework for testing gravitational phase coupling in quantum systems, as predicted by the Segmented Spacetime (SSZ) model. Building on Papers A and B, we perform a rigorous order-of-magnitude feasibility analysis revealing that the predicted SSZ effect at laboratory-scale height differences (Dh ~ mm) is approximately 12 orders of magnitude below the noise floor of current superconducting qubit technology. Crucially, a null result in this regime is itself SSZ-consistent: the theory predicts negligible effects at mm-scale with current coherence times. We therefore reframe this work as: (1) an upper-bound experiment that can constrain anomalous phase couplings in solid-state qubits, (2) a platform comparison identifying optical atomic clocks as the appropriate gold-standard test system (where DF ~ 0.3 rad at Dh = 1 m is readily detectable), and (3) a statistical falsification framework using slope-fitting rather than binary thresholds."""
    
    p = doc.add_paragraph(abstract)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run('Keywords: ').bold = True
    p.add_run('Gravitational Phase Coupling, Quantum Systems, Upper Bound, Feasibility Analysis, Optical Clocks, Falsifiability')
    
    doc.add_page_break()
    
    # 1. Introduction
    doc.add_heading('1. Introduction', level=1)
    
    doc.add_heading('1.1 Context from Papers A and B', level=2)
    p = doc.add_paragraph('In Papers A and B, we derived the SSZ prediction for gravitational phase drift. We showed this effect is deterministic, geometry-linked, and in principle compensable.')
    
    doc.add_heading('1.2 The Critical Question', level=2)
    p = doc.add_paragraph()
    p.add_run('Can this effect be detected with current technology?').bold = True
    
    p = doc.add_paragraph('This paper provides the honest answer: ')
    p.add_run('No, not at laboratory scales with superconducting qubits.').bold = True
    p.add_run(' However, this negative result is scientifically valuable--and is in fact ')
    p.add_run('consistent with SSZ predictions').bold = True
    p.add_run(' in the current regime.')
    
    doc.add_heading('1.3 Revised Goals', level=2)
    doc.add_paragraph('Quantify the feasibility gap between predicted signal and noise floor', style='List Number')
    doc.add_paragraph('Identify appropriate experimental platforms where detection is possible', style='List Number')
    doc.add_paragraph('Design an upper-bound experiment that provides value regardless of outcome', style='List Number')
    doc.add_paragraph('Establish a statistical framework for falsification claims', style='List Number')
    
    # 2. Feasibility Analysis
    doc.add_heading('2. Feasibility Analysis', level=1)
    
    doc.add_heading('2.1 Signal Size', level=2)
    p = doc.add_paragraph('For a 5 GHz qubit with Ramsey time T = 100 us:')
    
    table = doc.add_table(rows=5, cols=3)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    headers = ['Dh', 'DD_SSZ', 'DF (100 us)']
    for i, h in enumerate(headers):
        table.rows[0].cells[i].text = h
        table.rows[0].cells[i].paragraphs[0].runs[0].bold = True
        set_cell_shading(table.rows[0].cells[i], 'D9E2F3')
    
    data = [
        ('1 mm', '1.09x10^-19', '3.43x10^-13 rad'),
        ('1 m', '1.09x10^-16', '3.43x10^-10 rad'),
        ('10 m', '1.09x10^-15', '3.43x10^-9 rad'),
        ('100 m', '1.09x10^-14', '3.43x10^-8 rad'),
    ]
    for i, row in enumerate(data):
        for j, val in enumerate(row):
            table.rows[i+1].cells[j].text = val
    
    doc.add_paragraph()
    
    # PATCH 2: "representative" not "state-of-the-art"
    doc.add_heading('2.2 Noise Floor', level=2)
    p = doc.add_paragraph()
    p.add_run('Representative').bold = True
    p.add_run(' single-shot phase uncertainty in superconducting qubit measurements:')
    
    table2 = doc.add_table(rows=5, cols=3)
    table2.style = 'Table Grid'
    table2.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    table2.rows[0].cells[0].text = 'Source'
    table2.rows[0].cells[1].text = 'Magnitude'
    table2.rows[0].cells[2].text = 'Notes'
    for cell in table2.rows[0].cells:
        cell.paragraphs[0].runs[0].bold = True
        set_cell_shading(cell, 'D9E2F3')
    
    noise_data = [
        ('Quantum projection noise', 'O(1 rad)', 'Fundamental limit'),
        ('LO phase noise (100 us)', '~10^-3 rad', 'Oscillator dependent'),
        ('Temperature drift (1 mK)', '~0.6 rad', 'Frequency shift'),
        ('Combined', 'O(1 rad)', 'Dominated by projection'),
    ]
    for i, (src, mag, note) in enumerate(noise_data):
        table2.rows[i+1].cells[0].text = src
        table2.rows[i+1].cells[1].text = mag
        table2.rows[i+1].cells[2].text = note
    
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run('Note: ').italic = True
    p.add_run('These are representative order-of-magnitude estimates. Actual values depend on specific hardware and measurement protocols.')
    
    doc.add_heading('2.3 Averaging Requirements', level=2)
    
    table3 = doc.add_table(rows=4, cols=5)
    table3.style = 'Table Grid'
    table3.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    headers3 = ['Dh', 'Signal', 'N required', 'Time @ 10 kHz', 'Feasible?']
    for i, h in enumerate(headers3):
        table3.rows[0].cells[i].text = h
        table3.rows[0].cells[i].paragraphs[0].runs[0].bold = True
        set_cell_shading(table3.rows[0].cells[i], 'D9E2F3')
    
    avg_data = [
        ('1 mm', '3.4x10^-13 rad', '7.6x10^25', '2.4x10^14 years', 'No'),
        ('1 m', '3.4x10^-10 rad', '7.6x10^19', '2.4x10^8 years', 'No'),
        ('10 m', '3.4x10^-9 rad', '7.6x10^17', '2.4x10^6 years', 'No'),
    ]
    for i, row in enumerate(avg_data):
        for j, val in enumerate(row):
            table3.rows[i+1].cells[j].text = val
    
    doc.add_paragraph()
    
    doc.add_heading('2.4 Conclusion', level=2)
    p = doc.add_paragraph()
    p.add_run('The SSZ effect is ~12 orders of magnitude below detectability with current superconducting qubit technology. A null result is SSZ-consistent.').bold = True
    
    doc.add_page_break()
    
    # 3. Alternative Platforms - PATCH 1: Optical clock fix
    doc.add_heading('3. Alternative Platforms', level=1)
    
    doc.add_heading('3.1 Optical Atomic Clocks', level=2)
    
    p = doc.add_paragraph('Optical clocks operate at ~10^15 Hz with coherence times of seconds. The key advantage is the ')
    p.add_run('10^5x higher frequency').bold = True
    p.add_run(' combined with ')
    p.add_run('10^4x longer coherence').bold = True
    p.add_run(':')
    
    table4 = doc.add_table(rows=8, cols=4)
    table4.style = 'Table Grid'
    table4.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    headers4 = ['Parameter', 'Transmon Qubit', 'Optical Clock', 'Ratio']
    for i, h in enumerate(headers4):
        table4.rows[0].cells[i].text = h
        table4.rows[0].cells[i].paragraphs[0].runs[0].bold = True
        set_cell_shading(table4.rows[0].cells[i], 'D9E2F3')
    
    # CORRECTED VALUES
    platform_data = [
        ('Frequency f', '5 GHz', '429 THz', '8.6x10^4'),
        ('w = 2pf', '3.1x10^10 rad/s', '2.7x10^15 rad/s', '8.6x10^4'),
        ('Coherence t', '100 us', '1 s', '10^4'),
        ('DD_SSZ @ 1m', '1.09x10^-16', '1.09x10^-16', '1'),
        ('DF @ Dh=1m', '3.4x10^-10 rad', '~0.29 rad', '8.6x10^8'),
        ('N for SNR=3', '7.6x10^19', '~100', '--'),
        ('Feasible?', 'No', 'YES', '--'),
    ]
    for i, row in enumerate(platform_data):
        for j, val in enumerate(row):
            table4.rows[i+1].cells[j].text = val
    
    doc.add_paragraph()
    
    p = doc.add_paragraph()
    p.add_run('Calculation for optical clocks:').bold = True
    
    calc = doc.add_paragraph('DF = w x DD_SSZ x t = 2.7x10^15 rad/s x 1.09x10^-16 x 1 s = 0.29 rad')
    calc.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    p = doc.add_paragraph('This is a ')
    p.add_run('directly measurable').bold = True
    p.add_run(' phase shift. Optical clock experiments have already demonstrated gravitational redshift at the ~1 cm level (Bothwell et al., Nature 2022).')
    
    doc.add_heading('3.2 Recommendation', level=2)
    p = doc.add_paragraph()
    p.add_run('For testing SSZ predictions quantitatively, optical atomic clocks are the gold-standard platform.').bold = True
    p.add_run(' Superconducting qubits can provide upper bounds but cannot detect GR-level effects.')
    
    # 4. Upper-Bound Experiment
    doc.add_heading('4. Upper-Bound Experiment Design', level=1)
    
    doc.add_heading('4.1 Scientific Value', level=2)
    doc.add_paragraph('Constraining anomalous couplings: any beyond-GR coupling must be smaller than our bound', style='List Bullet')
    doc.add_paragraph('Validating null predictions: SSZ predicts negligibility at mm-scale--confirming this is a positive result', style='List Bullet')
    doc.add_paragraph('Establishing methodology: first systematic study of gravitational phase coupling in solid-state qubits', style='List Bullet')
    
    # PATCH 3: Chip tilt explanation
    doc.add_heading('4.2 Hardware Configurations', level=2)
    
    p = doc.add_paragraph()
    p.add_run('Configuration A: Chip Tilt').bold = True
    
    p = doc.add_paragraph('When a chip of length L is tilted by angle theta from horizontal, qubits at opposite ends experience a height difference:')
    
    eq = doc.add_paragraph('Dh = L x sin(theta)  (approximately L x theta for small angles)')
    eq.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    table5 = doc.add_table(rows=4, cols=3)
    table5.style = 'Table Grid'
    table5.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    table5.rows[0].cells[0].text = 'Tilt Angle'
    table5.rows[0].cells[1].text = 'sin(theta)'
    table5.rows[0].cells[2].text = 'Dh (20 mm chip)'
    for cell in table5.rows[0].cells:
        cell.paragraphs[0].runs[0].bold = True
        set_cell_shading(cell, 'D9E2F3')
    
    tilt_data = [('1 deg', '0.0175', '0.35 mm'), ('5 deg', '0.0872', '1.74 mm'), ('10 deg', '0.174', '3.47 mm')]
    for i, (angle, sin_t, dh) in enumerate(tilt_data):
        table5.rows[i+1].cells[0].text = angle
        table5.rows[i+1].cells[1].text = sin_t
        table5.rows[i+1].cells[2].text = dh
    
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run('Implementation: ').italic = True
    p.add_run('Precision goniometer stage under dilution refrigerator sample mount. Requires careful thermal management.')
    
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run('Configuration B: Remote Entanglement').bold = True
    p = doc.add_paragraph('Two qubits in separate dilution refrigerators at different heights (3-100 m), connected via microwave or optical link.')
    
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run('Configuration C: 3D Chiplet Stack').bold = True
    p = doc.add_paragraph('Vertically stacked quantum processors (0.5-5 mm Dh). Emerging technology pursued by IBM and Google.')
    
    doc.add_page_break()
    
    # 5. Statistical Framework - PATCH 4: Concrete example
    doc.add_heading('5. Statistical Falsification Framework', level=1)
    
    doc.add_heading('5.1 Model Comparison', level=2)
    
    p = doc.add_paragraph()
    p.add_run('M0 (Null): ').bold = True
    p.add_run('DF = 0 + noise')
    
    p = doc.add_paragraph()
    p.add_run('M_SSZ: ').bold = True
    p.add_run('DF = a_SSZ x Dh + noise (predicted slope)')
    
    p = doc.add_paragraph()
    p.add_run('M_anom: ').bold = True
    p.add_run('DF = a_fit x Dh + noise (free parameter)')
    
    doc.add_heading('5.2 Falsification Criteria', level=2)
    
    p = doc.add_paragraph()
    p.add_run('SSZ falsified if: ').bold = True
    p.add_run('measured slope inconsistent with prediction at >3s AND significantly non-zero')
    
    p = doc.add_paragraph()
    p.add_run('SSZ supported if: ').bold = True
    p.add_run('null result consistent with a_SSZ ~ 0 at mm-scale (this IS the prediction)')
    
    doc.add_heading('5.3 Upper Bound: Concrete Example', level=2)
    
    p = doc.add_paragraph('With Dh_max = 3.5 mm (10 deg tilt), N = 10^9 shots, s_single ~ 1 rad:')
    
    doc.add_paragraph('s_after_avg = s_single / sqrt(N) = 1 / sqrt(10^9) = 3.2x10^-5 rad', style='List Bullet')
    doc.add_paragraph('s_slope = s_after_avg / Dh_max = 3.2x10^-5 / 3.5x10^-3 = 9x10^-3 rad/m', style='List Bullet')
    doc.add_paragraph('Upper bound: |a_anom| < 9x10^-3 rad/m (95% CL)', style='List Bullet')
    
    p = doc.add_paragraph('For comparison, SSZ-predicted slope: a_SSZ ~ 6.7x10^-13 rad/m')
    
    p = doc.add_paragraph()
    p.add_run('This experiment constrains anomalous couplings to < 10^10 x a_SSZ').bold = True
    p.add_run('--scientifically meaningful as a first systematic bound.')
    
    # 6. Confounds
    doc.add_heading('6. Confound Discrimination', level=1)
    
    doc.add_heading('6.1 Scaling Signatures (not absolute exclusions)', level=2)
    
    table6 = doc.add_table(rows=6, cols=4)
    table6.style = 'Table Grid'
    table6.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    headers6 = ['Source', 'Dh scaling', 'w scaling', 't scaling']
    for i, h in enumerate(headers6):
        table6.rows[0].cells[i].text = h
        table6.rows[0].cells[i].paragraphs[0].runs[0].bold = True
        set_cell_shading(table6.rows[0].cells[i], 'D9E2F3')
    
    confound_data = [
        ('SSZ', 'Linear', 'Linear', 'Linear'),
        ('Temperature', 'May correlate', 'Weak', 'Non-linear'),
        ('LO noise', 'None', 'None', 'sqrt(t)'),
        ('Vibration', 'Mechanical', 'None', 'AC spectrum'),
        ('EM crosstalk', 'Position-dep.', 'Weak', 'Constant'),
    ]
    for i, row in enumerate(confound_data):
        for j, val in enumerate(row):
            table6.rows[i+1].cells[j].text = val
    
    doc.add_paragraph()
    
    doc.add_heading('6.2 Key Controls', level=2)
    doc.add_paragraph('Randomize Dh order to break thermal correlation', style='List Bullet')
    doc.add_paragraph('Reference qubits for common-mode subtraction', style='List Bullet')
    doc.add_paragraph('Accelerometer monitoring for vibration correlation', style='List Bullet')
    
    # 7. Consistency
    doc.add_heading('7. Consistency with Papers A and B', level=1)
    
    p = doc.add_paragraph()
    p.add_run('Resolution: ').bold = True
    p.add_run('Papers A/B describe the regime where SSZ becomes relevant (future systems with T2 >> 1s, Dh ~ m). Paper C tests current systems where SSZ is negligible. ')
    p.add_run('A null result today validates SSZ in the regime where it predicts negligibility.').bold = True
    
    # 8. Conclusion
    doc.add_heading('8. Conclusion', level=1)
    
    doc.add_paragraph('Feasibility: SSZ effect is ~12 orders of magnitude below current sensitivity', style='List Number')
    doc.add_paragraph('Null is positive: null result in current regime is SSZ-consistent', style='List Number')
    doc.add_paragraph('Platform: optical clocks are gold-standard (DF ~ 0.3 rad at 1m)', style='List Number')
    doc.add_paragraph('Statistics: slope-fitting with explicit confidence levels', style='List Number')
    doc.add_paragraph('Value: first systematic constraint on gravitational phase coupling in qubits', style='List Number')
    
    # References
    doc.add_heading('References', level=1)
    doc.add_paragraph('[1] Casu & Wrede (2025). Paper A: Geometric Qubit Optimization.')
    doc.add_paragraph('[2] Casu & Wrede (2025). Paper B: Phase Coherence and Entanglement.')
    doc.add_paragraph('[3] Bothwell et al. (2022). Nature 602, 420-424.')
    doc.add_paragraph('[4] https://github.com/error-wtf/ssz-qubits')
    
    # Footer
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run('(c) 2025 Carmen Wrede & Lino Casu | ANTI-CAPITALIST SOFTWARE LICENSE v1.4').italic = True
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Save
    os.makedirs(OUTPUT_DIR, exist_ok=True)
    os.makedirs(PAPERS_DIR, exist_ok=True)
    
    path1 = os.path.join(PAPERS_DIR, 'SSZ_Paper_C_v1.2_Bulletproof.docx')
    path2 = os.path.join(OUTPUT_DIR, 'SSZ_Paper_C_v1.2_Bulletproof.docx')
    
    doc.save(path1)
    doc.save(path2)
    
    print(f"Saved: {path1}")
    print(f"Saved: {path2}")
    
    return path1

if __name__ == "__main__":
    create_paper_c_v12()
