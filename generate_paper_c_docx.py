#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Generate Paper C DOCX Document

Creates a professional Word document for Paper C:
"Falsifiable Predictions and Experimental Protocols for SSZ-Aware Quantum Computing"

Consistent with Paper A and Paper B formatting.

(c) 2025 Carmen Wrede, Lino Casu
"""

import os
import sys
from docx import Document
from docx.shared import Inches, Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# Paths
SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))
OUTPUT_DIR = os.path.join(SCRIPT_DIR, 'outputs')
PAPERS_DIR = r'E:\clone\SSZ_QUBIT_PAPERS'

def set_cell_shading(cell, color):
    """Set cell background color."""
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color)
    cell._tc.get_or_add_tcPr().append(shading)

def add_heading_with_number(doc, text, level=1):
    """Add a numbered heading."""
    heading = doc.add_heading(text, level=level)
    heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
    return heading

def create_paper_c_docx():
    """Create the Paper C DOCX document."""
    doc = Document()
    
    # Set up styles
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(11)
    
    # Title
    title = doc.add_paragraph()
    title_run = title.add_run('Falsifiable Predictions and Experimental Protocols\nfor SSZ-Aware Quantum Computing')
    title_run.bold = True
    title_run.font.size = Pt(16)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Subtitle
    subtitle = doc.add_paragraph()
    subtitle_run = subtitle.add_run('Paper C: Experimental Validation Framework')
    subtitle_run.italic = True
    subtitle_run.font.size = Pt(12)
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()
    
    # Authors
    authors = doc.add_paragraph()
    authors_run = authors.add_run('Lino Casu, Carmen Wrede')
    authors_run.bold = True
    authors.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    affiliation = doc.add_paragraph()
    affiliation.add_run('Independent Researchers').italic = True
    affiliation.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    contact = doc.add_paragraph()
    contact.add_run('Contact: mail@error.wtf')
    contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    date = doc.add_paragraph()
    date.add_run('December 2025')
    date.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()
    
    # =========================================================================
    # Abstract
    # =========================================================================
    doc.add_heading('Abstract', level=1)
    
    abstract_text = """We present a comprehensive framework for experimentally testing Segmented Spacetime (SSZ) effects in quantum computing systems. Building on our previous work demonstrating geometric qubit optimization (Paper A) and phase coherence preservation (Paper B), this paper provides: (1) five quantitative, falsifiable predictions derived from SSZ theory, (2) a detailed experimental protocol for the gold-standard falsification test, (3) strategies for discriminating SSZ effects from common confounds, and (4) realistic measurement requirements. Our analysis shows that SSZ predictions are testable with current superconducting qubit technology using piezo-controlled height stages, Ramsey interferometry, and differential measurement techniques. The key discriminator is the with/without compensation test: SSZ predicts a deterministic, geometry-linked phase drift that can be cancelled by applying the predicted correction, while confounds cannot be compensated geometrically. We provide all numerical predictions with explicit falsification thresholds."""
    
    abstract = doc.add_paragraph(abstract_text)
    abstract.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
    # Keywords
    keywords = doc.add_paragraph()
    keywords.add_run('Keywords: ').bold = True
    keywords.add_run('Segmented Spacetime, Quantum Computing, Falsifiability, Experimental Protocol, Phase Drift, Gravitational Effects')
    
    doc.add_page_break()
    
    # =========================================================================
    # 1. Introduction
    # =========================================================================
    doc.add_heading('1. Introduction', level=1)
    
    doc.add_heading('1.1 Motivation', level=2)
    p = doc.add_paragraph()
    p.add_run('The Segmented Spacetime (SSZ) framework predicts that gravitational time dilation introduces deterministic phase biases in quantum systems. In Papers A and B, we demonstrated:')
    
    bullet1 = doc.add_paragraph('Paper A: Height differences as small as 1 mm create measurable segment density gradients that affect gate timing', style='List Bullet')
    bullet2 = doc.add_paragraph('Paper B: These effects accumulate coherently over time, enabling compensation via calibrated phase corrections', style='List Bullet')
    
    p = doc.add_paragraph()
    p.add_run('A critical question remains: ').italic = True
    p.add_run('How can these predictions be experimentally tested and potentially falsified?').bold = True
    
    doc.add_heading('1.2 The Falsifiability Criterion', level=2)
    p = doc.add_paragraph('For SSZ to be a legitimate scientific theory, it must make predictions that are:')
    doc.add_paragraph('Quantitative — specific numerical values, not just qualitative trends', style='List Number')
    doc.add_paragraph('Testable — measurable with available technology', style='List Number')
    doc.add_paragraph('Falsifiable — there must be outcomes that would disprove the theory', style='List Number')
    
    p = doc.add_paragraph('This paper addresses all three criteria explicitly.')
    
    doc.add_heading('1.3 Summary of Key Results', level=2)
    
    # Summary table
    table = doc.add_table(rows=6, cols=4)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    headers = ['Prediction', 'SSZ Value', 'Falsification Threshold', 'Testable?']
    header_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        header_cells[i].text = header
        header_cells[i].paragraphs[0].runs[0].bold = True
        set_cell_shading(header_cells[i], 'D9E2F3')
    
    data = [
        ('Phase drift at Δh=1mm', '3.43×10⁻¹⁵ rad/μs', '<1.72×10⁻¹⁵ rad/μs', 'Yes'),
        ('Coherent zone (ε=10⁻¹⁸)', '18.3 mm', '<9.2 mm', 'Yes'),
        ('ω-scaling (7GHz/5GHz)', '1.40', '<1.20', 'Yes'),
        ('Compensation efficiency', '99%', '<90%', 'Yes'),
        ('Cross-zone drift/gate', '6.28×10⁻¹⁵ rad', '<3.14×10⁻¹⁵ rad', 'Yes'),
    ]
    
    for i, row_data in enumerate(data):
        row_cells = table.rows[i + 1].cells
        for j, cell_data in enumerate(row_data):
            row_cells[j].text = cell_data
    
    doc.add_paragraph()
    
    # =========================================================================
    # 2. Theoretical Foundation
    # =========================================================================
    doc.add_heading('2. Theoretical Foundation', level=1)
    
    doc.add_heading('2.1 Core SSZ Equations', level=2)
    p = doc.add_paragraph('The SSZ time dilation factor at radius r from mass M is:')
    
    eq1 = doc.add_paragraph()
    eq1.add_run('D').italic = True
    eq1.add_run('SSZ').font.subscript = True
    eq1.add_run('(r) = 1 / (1 + Ξ(r))')
    eq1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    p = doc.add_paragraph('where the segment density Ξ(r) in the weak-field regime (r >> r')
    p.add_run('s').font.subscript = True
    p.add_run(') is:')
    
    eq2 = doc.add_paragraph()
    eq2.add_run('Ξ(r) = r').italic = True
    eq2.add_run('s').font.subscript = True
    eq2.add_run(' / (2r)')
    eq2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    p = doc.add_paragraph('with Schwarzschild radius r')
    p.add_run('s').font.subscript = True
    p.add_run(' = 2GM/c².')
    
    doc.add_heading('2.2 Phase Drift Formula', level=2)
    p = doc.add_paragraph('For two qubits at heights h₁ and h₂ (with Δh = h₂ - h₁), the differential time dilation produces a phase drift:')
    
    eq3 = doc.add_paragraph()
    eq3.add_run('ΔΦ(t) = ω × ΔD').italic = True
    eq3.add_run('SSZ').font.subscript = True
    eq3.add_run('(Δh) × t')
    eq3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    p = doc.add_paragraph('where ω = 2πf is the qubit angular frequency.')
    
    doc.add_heading('2.3 Segment-Coherent Zones', level=2)
    p = doc.add_paragraph('Qubits within a "coherent zone" share timing to within tolerance ε. The zone width is:')
    
    eq4 = doc.add_paragraph()
    eq4.add_run('z(ε) = 4 × ε × R² / r').italic = True
    eq4.add_run('s').font.subscript = True
    eq4.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Zone width table
    p = doc.add_paragraph('For Earth (R = 6.371×10⁶ m, r')
    p.add_run('s').font.subscript = True
    p.add_run(' = 8.87×10⁻³ m):')
    
    table2 = doc.add_table(rows=6, cols=2)
    table2.style = 'Table Grid'
    table2.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    table2.rows[0].cells[0].text = 'Tolerance ε'
    table2.rows[0].cells[1].text = 'Zone Width z'
    table2.rows[0].cells[0].paragraphs[0].runs[0].bold = True
    table2.rows[0].cells[1].paragraphs[0].runs[0].bold = True
    set_cell_shading(table2.rows[0].cells[0], 'D9E2F3')
    set_cell_shading(table2.rows[0].cells[1], 'D9E2F3')
    
    zone_data = [
        ('10⁻¹⁶', '1.83 m'),
        ('10⁻¹⁷', '183 mm'),
        ('10⁻¹⁸', '18.3 mm'),
        ('10⁻¹⁹', '1.83 mm'),
        ('10⁻²⁰', '183 μm'),
    ]
    
    for i, (eps, width) in enumerate(zone_data):
        table2.rows[i + 1].cells[0].text = eps
        table2.rows[i + 1].cells[1].text = width
    
    doc.add_paragraph()
    
    # =========================================================================
    # 3. Falsifiable Predictions
    # =========================================================================
    doc.add_heading('3. Falsifiable Predictions', level=1)
    
    # Figure 1 placeholder
    p = doc.add_paragraph()
    p.add_run('[Figure 1: SSZ Phase Drift vs Height Difference - see outputs/paper_c_fig1_phase_vs_height.png]').italic = True
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Try to add actual figure
    fig1_path = os.path.join(OUTPUT_DIR, 'paper_c_fig1_phase_vs_height.png')
    if os.path.exists(fig1_path):
        doc.add_picture(fig1_path, width=Inches(5.5))
        last_paragraph = doc.paragraphs[-1]
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        caption = doc.add_paragraph('Figure 1: SSZ phase drift rate as a function of height difference (log-log scale).')
        caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
        caption.runs[0].italic = True
        caption.runs[0].font.size = Pt(10)
    
    doc.add_paragraph()
    
    doc.add_heading('3.1 Prediction 1: Phase Drift Rate at Δh = 1 mm', level=2)
    
    p = doc.add_paragraph()
    p.add_run('Formula: ').bold = True
    p.add_run('ΔΦ/t = ω × ΔD')
    p.add_run('SSZ').font.subscript = True
    p.add_run('(Δh)')
    
    p = doc.add_paragraph()
    p.add_run('Numerical Prediction:').bold = True
    doc.add_paragraph('Qubit frequency: f = 5 GHz', style='List Bullet')
    doc.add_paragraph('Height difference: Δh = 1 mm', style='List Bullet')
    doc.add_paragraph('Predicted phase drift rate: 3.43 × 10⁻¹⁵ rad/μs', style='List Bullet')
    
    p = doc.add_paragraph()
    p.add_run('Falsification Threshold: ').bold = True
    p.add_run('If measured drift is < 1.72 × 10⁻¹⁵ rad/μs (less than 50% of prediction), SSZ is falsified.')
    
    doc.add_heading('3.2 Prediction 2: Coherent Zone Width at ε = 10⁻¹⁸', level=2)
    
    p = doc.add_paragraph()
    p.add_run('Formula: ').bold = True
    p.add_run('z(ε) = 4εR²/r')
    p.add_run('s').font.subscript = True
    
    p = doc.add_paragraph()
    p.add_run('Predicted zone width: 18.3 mm').bold = True
    
    p = doc.add_paragraph()
    p.add_run('Falsification Threshold: ').bold = True
    p.add_run('If measured zone is < 9.2 mm, SSZ is falsified.')
    
    doc.add_heading('3.3 Prediction 3: Frequency Scaling Ratio', level=2)
    
    p = doc.add_paragraph()
    p.add_run('Formula: ').bold = True
    p.add_run('ΔΦ(ω₂)/ΔΦ(ω₁) = ω₂/ω₁')
    
    p = doc.add_paragraph()
    p.add_run('Predicted ratio (7 GHz vs 5 GHz): 1.40').bold = True
    
    p = doc.add_paragraph()
    p.add_run('Falsification Threshold: ').bold = True
    p.add_run('If measured ratio is < 1.20 or > 1.60, SSZ is falsified.')
    
    doc.add_heading('3.4 Prediction 4: Compensation Efficiency', level=2)
    
    p = doc.add_paragraph()
    p.add_run('Formula: ').bold = True
    p.add_run('η = 1 - |ΔΦ')
    p.add_run('compensated').font.subscript = True
    p.add_run('|/|ΔΦ')
    p.add_run('uncompensated').font.subscript = True
    p.add_run('|')
    
    p = doc.add_paragraph()
    p.add_run('Predicted efficiency: η ≥ 99%').bold = True
    
    p = doc.add_paragraph()
    p.add_run('Falsification Threshold: ').bold = True
    p.add_run('If compensation achieves < 90% reduction, SSZ is falsified.')
    
    doc.add_heading('3.5 Prediction 5: Cross-Zone Phase Drift per Gate', level=2)
    
    p = doc.add_paragraph()
    p.add_run('Formula: ').bold = True
    p.add_run('ΔΦ')
    p.add_run('gate').font.subscript = True
    p.add_run(' = ω × ΔXi × t')
    p.add_run('gate').font.subscript = True
    
    p = doc.add_paragraph()
    p.add_run('Predicted drift (at 2× zone boundary): 6.28 × 10⁻¹⁵ rad/gate').bold = True
    
    p = doc.add_paragraph()
    p.add_run('Falsification Threshold: ').bold = True
    p.add_run('If measured drift is < 3.14 × 10⁻¹⁵ rad/gate, SSZ is falsified.')
    
    doc.add_page_break()
    
    # =========================================================================
    # 4. Gold-Standard Falsification Experiment
    # =========================================================================
    doc.add_heading('4. Gold-Standard Falsification Experiment', level=1)
    
    doc.add_heading('4.1 Experimental Setup', level=2)
    
    table3 = doc.add_table(rows=9, cols=2)
    table3.style = 'Table Grid'
    table3.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    table3.rows[0].cells[0].text = 'Component'
    table3.rows[0].cells[1].text = 'Specification'
    table3.rows[0].cells[0].paragraphs[0].runs[0].bold = True
    table3.rows[0].cells[1].paragraphs[0].runs[0].bold = True
    set_cell_shading(table3.rows[0].cells[0], 'D9E2F3')
    set_cell_shading(table3.rows[0].cells[1], 'D9E2F3')
    
    setup_data = [
        ('Qubit type', 'Superconducting transmon'),
        ('Frequency', '5 GHz (primary), 7 GHz (scaling test)'),
        ('T₂ time', '≥ 100 μs'),
        ('Height control', 'Piezo translation stage'),
        ('Height range', '0.1 mm to 10 mm'),
        ('Height precision', '1 μm'),
        ('Temperature stability', '< 1 mK'),
        ('Vibration isolation', 'Active + passive'),
    ]
    
    for i, (comp, spec) in enumerate(setup_data):
        table3.rows[i + 1].cells[0].text = comp
        table3.rows[i + 1].cells[1].text = spec
    
    doc.add_paragraph()
    
    doc.add_heading('4.2 Protocol', level=2)
    
    p = doc.add_paragraph()
    p.add_run('CALIBRATION PHASE').bold = True
    doc.add_paragraph('Calibrate qubits at reference height (h = 0)', style='List Bullet')
    doc.add_paragraph('Characterize T₁, T₂, gate fidelities', style='List Bullet')
    doc.add_paragraph('Record temperature, vibration baseline', style='List Bullet')
    
    p = doc.add_paragraph()
    p.add_run('DATA ACQUISITION PHASE').bold = True
    doc.add_paragraph('Set height difference Δh using piezo stage', style='List Bullet')
    doc.add_paragraph('Prepare Bell state |Φ⁺⟩ = (|00⟩ + |11⟩)/√2', style='List Bullet')
    doc.add_paragraph('Perform Ramsey sequence to measure phase ΔΦ(Δh, t)', style='List Bullet')
    doc.add_paragraph('Record Bell state fidelity F(Δh, t)', style='List Bullet')
    doc.add_paragraph('Repeat with SSZ compensation enabled', style='List Bullet')
    doc.add_paragraph('Randomize Δh order to avoid systematic drift', style='List Bullet')
    
    p = doc.add_paragraph()
    p.add_run('CONTROL MEASUREMENTS').bold = True
    doc.add_paragraph('Null test: same Δh, different thermal gradient', style='List Bullet')
    doc.add_paragraph('Position test: same height, different x-y position', style='List Bullet')
    doc.add_paragraph('Frequency test: repeat at 7 GHz', style='List Bullet')
    
    doc.add_heading('4.3 Expected Results', level=2)
    
    # Figure 4 placeholder
    fig4_path = os.path.join(OUTPUT_DIR, 'paper_c_fig4_falsification.png')
    if os.path.exists(fig4_path):
        doc.add_picture(fig4_path, width=Inches(6))
        last_paragraph = doc.paragraphs[-1]
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        caption = doc.add_paragraph('Figure 4: Expected results from the gold-standard falsification experiment.')
        caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
        caption.runs[0].italic = True
        caption.runs[0].font.size = Pt(10)
    
    doc.add_paragraph()
    
    doc.add_heading('4.4 Signatures of Success', level=2)
    
    doc.add_paragraph('ΔΦ monotonically increases with Δh — linear on log-log plot with slope = 1', style='List Number')
    doc.add_paragraph('ΔΦ scales linearly with t — same slope across all Δh', style='List Number')
    doc.add_paragraph('ΔΦ scales linearly with ω — 7GHz/5GHz ratio = 1.40 ± 0.05', style='List Number')
    doc.add_paragraph('Compensation removes >90% of ΔΦ — differential test passes', style='List Number')
    doc.add_paragraph('Effect is reproducible — multiple runs give consistent results', style='List Number')
    
    doc.add_page_break()
    
    # =========================================================================
    # 5. Confound Discrimination
    # =========================================================================
    doc.add_heading('5. Confound Discrimination', level=1)
    
    # Figure 5 - Confound Matrix
    fig5_path = os.path.join(OUTPUT_DIR, 'paper_c_fig5_confounds.png')
    if os.path.exists(fig5_path):
        doc.add_picture(fig5_path, width=Inches(5))
        last_paragraph = doc.paragraphs[-1]
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        caption = doc.add_paragraph('Figure 5: Confound discrimination matrix. Only SSZ passes all five discriminating tests.')
        caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
        caption.runs[0].italic = True
        caption.runs[0].font.size = Pt(10)
    
    doc.add_paragraph()
    
    doc.add_heading('5.1 Temperature Drift', level=2)
    p = doc.add_paragraph()
    p.add_run('Effect: ').bold = True
    p.add_run('Qubit frequency shifts ~1 kHz/mK → phase drift')
    
    p = doc.add_paragraph()
    p.add_run('Why it\'s not SSZ:').bold = True
    doc.add_paragraph('Temperature effect is NOT monotonic in Δh', style='List Bullet')
    doc.add_paragraph('Temperature effect does NOT scale with ω', style='List Bullet')
    doc.add_paragraph('SSZ compensation would NOT reduce temperature-induced drift', style='List Bullet')
    
    doc.add_heading('5.2 LO Phase Noise', level=2)
    p = doc.add_paragraph()
    p.add_run('Effect: ').bold = True
    p.add_run('Random phase fluctuations ~10⁻³ rad/√Hz')
    
    p = doc.add_paragraph()
    p.add_run('Why it\'s not SSZ:').bold = True
    doc.add_paragraph('LO noise is STOCHASTIC; SSZ is DETERMINISTIC', style='List Bullet')
    doc.add_paragraph('SSZ can be COMPENSATED; noise cannot', style='List Bullet')
    
    doc.add_heading('5.3 Mechanical Vibration', level=2)
    p = doc.add_paragraph()
    p.add_run('Effect: ').bold = True
    p.add_run('Height fluctuations → apparent phase noise')
    
    p = doc.add_paragraph()
    p.add_run('Why it\'s not SSZ:').bold = True
    doc.add_paragraph('Vibration is AC (time-varying); SSZ is DC (constant at fixed Δh)', style='List Bullet')
    doc.add_paragraph('SSZ prediction is for static Δh', style='List Bullet')
    
    doc.add_heading('5.4 The Key Discriminator: Compensation Test', level=2)
    
    # Figure 3 - Compensation
    fig3_path = os.path.join(OUTPUT_DIR, 'paper_c_fig3_compensation.png')
    if os.path.exists(fig3_path):
        doc.add_picture(fig3_path, width=Inches(6))
        last_paragraph = doc.paragraphs[-1]
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        caption = doc.add_paragraph('Figure 3: Compensation efficiency. (a) Fidelity vs gate count, (b) phase drift with/without compensation.')
        caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
        caption.runs[0].italic = True
        caption.runs[0].font.size = Pt(10)
    
    doc.add_paragraph()
    
    p = doc.add_paragraph('SSZ predicts a ')
    p.add_run('deterministic, geometry-linked').bold = True
    p.add_run(' phase drift. If we apply the predicted correction -ΔΦ')
    p.add_run('SSZ').font.subscript = True
    p.add_run(':')
    
    doc.add_paragraph('SSZ contribution is cancelled', style='List Bullet')
    doc.add_paragraph('Confound contributions are UNCHANGED', style='List Bullet')
    
    p = doc.add_paragraph()
    p.add_run('This differential test is powerful because:').bold = True
    doc.add_paragraph('It does not require knowing the absolute SSZ magnitude', style='List Number')
    doc.add_paragraph('It tests the correctability of the effect', style='List Number')
    doc.add_paragraph('Confounds cannot be "compensated" by geometry-based corrections', style='List Number')
    
    doc.add_page_break()
    
    # =========================================================================
    # 6. Measurement Requirements
    # =========================================================================
    doc.add_heading('6. Measurement Requirements', level=1)
    
    doc.add_heading('6.1 Phase Precision', level=2)
    doc.add_paragraph('SSZ signal at Δh = 1mm: 3.43 × 10⁻¹⁵ rad/μs', style='List Bullet')
    doc.add_paragraph('Required precision: 3.4 × 10⁻¹⁶ rad/μs (for SNR > 10)', style='List Bullet')
    doc.add_paragraph('Achievability: Standard Ramsey with ~10⁶ averages over T₂', style='List Bullet')
    
    doc.add_heading('6.2 Height Precision', level=2)
    doc.add_paragraph('Required: ~100 μm for 10% phase precision', style='List Bullet')
    doc.add_paragraph('Achievability: Piezo stages routinely achieve < 1 μm — easily met', style='List Bullet')
    
    doc.add_heading('6.3 Temperature Stability', level=2)
    doc.add_paragraph('For absolute measurement: ~10⁻¹² K (unrealistic)', style='List Bullet')
    doc.add_paragraph('For differential measurement: ~1 mK (achievable)', style='List Bullet')
    doc.add_paragraph('Solution: Use differential protocol to cancel common-mode temperature drift', style='List Bullet')
    
    # =========================================================================
    # 7. Discussion
    # =========================================================================
    doc.add_heading('7. Discussion', level=1)
    
    doc.add_heading('7.1 Why This Experiment is Decisive', level=2)
    doc.add_paragraph('Multiple independent signatures — All five predictions must hold simultaneously', style='List Number')
    doc.add_paragraph('Scaling tests — Wrong scaling immediately falsifies SSZ', style='List Number')
    doc.add_paragraph('Compensation test — Direct test of the core SSZ prediction', style='List Number')
    doc.add_paragraph('Confound discrimination — Each confound has distinct signatures', style='List Number')
    doc.add_paragraph('Reproducibility requirement — Effect must be consistent across runs', style='List Number')
    
    doc.add_heading('7.2 What Falsification Would Mean', level=2)
    p = doc.add_paragraph('If the experiment shows no Δh-dependence, wrong ω/t scaling, or compensation failure, then SSZ as an operational model for quantum systems is ')
    p.add_run('falsified').bold = True
    p.add_run('.')
    
    doc.add_heading('7.3 What Confirmation Would Mean', level=2)
    p = doc.add_paragraph('If all predictions are confirmed:')
    doc.add_paragraph('SSZ provides a valid operational model for gravitational effects in quantum systems', style='List Bullet')
    doc.add_paragraph('Practical benefit: Calibration and compensation strategies for quantum hardware', style='List Bullet')
    doc.add_paragraph('Theoretical interest: Supports the segment-based interpretation of spacetime', style='List Bullet')
    
    doc.add_heading('7.4 Relation to Papers A and B', level=2)
    
    table4 = doc.add_table(rows=4, cols=3)
    table4.style = 'Table Grid'
    table4.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    table4.rows[0].cells[0].text = 'Paper'
    table4.rows[0].cells[1].text = 'Focus'
    table4.rows[0].cells[2].text = 'Key Result'
    for cell in table4.rows[0].cells:
        cell.paragraphs[0].runs[0].bold = True
        set_cell_shading(cell, 'D9E2F3')
    
    table4.rows[1].cells[0].text = 'A'
    table4.rows[1].cells[1].text = 'Geometric optimization'
    table4.rows[1].cells[2].text = 'Gate timing correction ~10% fidelity improvement'
    
    table4.rows[2].cells[0].text = 'B'
    table4.rows[2].cells[1].text = 'Phase coherence'
    table4.rows[2].cells[2].text = 'Entanglement duration extended via SSZ compensation'
    
    table4.rows[3].cells[0].text = 'C'
    table4.rows[3].cells[1].text = 'Falsifiability'
    table4.rows[3].cells[2].text = '5 quantitative predictions with explicit thresholds'
    
    doc.add_paragraph()
    
    # =========================================================================
    # 8. Conclusion
    # =========================================================================
    doc.add_heading('8. Conclusion', level=1)
    
    p = doc.add_paragraph('We have presented a complete framework for experimentally testing SSZ predictions in quantum computing systems:')
    
    doc.add_paragraph('Five falsifiable predictions with explicit numerical values and thresholds', style='List Number')
    doc.add_paragraph('A detailed experimental protocol using current technology', style='List Number')
    doc.add_paragraph('Confound discrimination strategies for each major systematic effect', style='List Number')
    doc.add_paragraph('Realistic measurement requirements that are achievable with state-of-the-art hardware', style='List Number')
    
    p = doc.add_paragraph('The key insight is that SSZ predicts a ')
    p.add_run('deterministic, geometry-linked, compensable').bold = True
    p.add_run(' phase drift. The with/without compensation differential test is the strongest discriminator because confounds cannot be geometrically compensated.')
    
    p = doc.add_paragraph()
    p.add_run('We encourage experimental groups with access to superconducting qubit systems and precision height control to perform this test. The SSZ framework makes bold predictions — it deserves an equally rigorous experimental examination.').italic = True
    
    # =========================================================================
    # Acknowledgments
    # =========================================================================
    doc.add_heading('Acknowledgments', level=1)
    doc.add_paragraph('We thank the open-source community for valuable discussions and feedback.')
    
    # =========================================================================
    # References
    # =========================================================================
    doc.add_heading('References', level=1)
    
    doc.add_paragraph('[1] Casu, L. & Wrede, C. (2025). Paper A: Geometric Qubit Optimization via Segmented Spacetime. ssz-qubits repository.')
    doc.add_paragraph('[2] Casu, L. & Wrede, C. (2025). Paper B: Phase Coherence and Entanglement Preservation via Spacetime Segmentation. ssz-qubits repository.')
    doc.add_paragraph('[3] SSZ-Qubits Repository: https://github.com/error-wtf/ssz-qubits')
    
    # =========================================================================
    # Footer
    # =========================================================================
    doc.add_paragraph()
    footer = doc.add_paragraph()
    footer.add_run('© 2025 Carmen Wrede & Lino Casu').italic = True
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    license_p = doc.add_paragraph()
    license_p.add_run('Licensed under the ANTI-CAPITALIST SOFTWARE LICENSE v1.4').italic = True
    license_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Save
    output_path = os.path.join(PAPERS_DIR, 'SSZ_Paper_C_Falsifiability_Protocol_FINAL.docx')
    doc.save(output_path)
    print(f"Saved: {output_path}")
    
    # Also save to ssz-qubits outputs
    output_path2 = os.path.join(OUTPUT_DIR, 'SSZ_Paper_C_Falsifiability_Protocol.docx')
    doc.save(output_path2)
    print(f"Saved: {output_path2}")
    
    return output_path


if __name__ == "__main__":
    create_paper_c_docx()
