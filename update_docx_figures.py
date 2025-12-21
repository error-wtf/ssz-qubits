#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Update DOCX files with new supplementary figures.
Adds the 6 new plots to SSZ_Paper_D_MASTER.docx

(c) 2025 Carmen Wrede, Lino Casu
"""

import os
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

OUTPUT_DIR = os.path.join(os.path.dirname(__file__), 'outputs')

# New figures to add
NEW_FIGURES = [
    ('fig_strong_field_ssz_vs_gr.png', 'Figure S.1: Strong Field Comparison - D_SSZ vs D_GR at horizon. SSZ remains finite (0.555) while GR diverges to zero.'),
    ('fig_validation_summary.png', 'Figure S.2: Validation Summary - GPS, Pound-Rebka, NIST, and Tokyo Skytree experiments all match SSZ predictions within error bars.'),
    ('fig_phi_geometry.png', 'Figure S.3: Golden Ratio Geometry - Why phi controls the saturation rate in SSZ segment density.'),
    ('fig_coherent_zone_scaling.png', 'Figure S.4: Coherent Zone Scaling - Zone width vs tolerance epsilon for qubit placement optimization.'),
    ('fig_time_drift_accumulation.png', 'Figure S.5: Time Drift Accumulation - Phase drift growth over integration time and gate operations.'),
    ('fig_qubit_height_sensitivity.png', 'Figure S.6: Qubit Height Sensitivity - Segment density changes and phase drift for different height configurations.'),
]


def add_figures_to_docx(docx_path, output_path=None):
    """Add new supplementary figures to a DOCX file."""
    if output_path is None:
        base, ext = os.path.splitext(docx_path)
        output_path = f"{base}_with_supplementary{ext}"
    
    print(f"Loading: {docx_path}")
    doc = Document(docx_path)
    
    # Add section header
    doc.add_page_break()
    heading = doc.add_heading('Supplementary Figures', level=1)
    
    # Add each figure
    for filename, caption in NEW_FIGURES:
        filepath = os.path.join(OUTPUT_DIR, filename)
        
        if not os.path.exists(filepath):
            print(f"  [SKIP] {filename} not found")
            continue
        
        print(f"  [ADD] {filename}")
        
        # Add figure
        doc.add_picture(filepath, width=Inches(6.0))
        
        # Center the image
        last_paragraph = doc.paragraphs[-1]
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add caption
        caption_para = doc.add_paragraph()
        caption_run = caption_para.add_run(caption)
        caption_run.italic = True
        caption_run.font.size = Pt(10)
        caption_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add spacing
        doc.add_paragraph()
    
    # Save
    doc.save(output_path)
    print(f"\nSaved: {output_path}")
    return output_path


def main():
    print("=" * 60)
    print("SSZ DOCX Figure Updater")
    print("=" * 60)
    
    # Update Paper D Master
    paper_d = os.path.join(OUTPUT_DIR, 'SSZ_Paper_D_MASTER.docx')
    if os.path.exists(paper_d):
        add_figures_to_docx(paper_d)
    else:
        print(f"[ERROR] Not found: {paper_d}")
    
    print("\n" + "=" * 60)
    print("Done!")
    print("=" * 60)


if __name__ == "__main__":
    main()
